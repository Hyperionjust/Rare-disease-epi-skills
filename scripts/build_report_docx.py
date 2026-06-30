#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_report_docx.py  --  Word (.docx) report generator for the `rare-disease-epi` skill.

    python scripts/build_report_docx.py <data.json> -o <output.docx>

Reads the SAME structured JSON as build_report.py (schema: templates/report_schema.json)
and produces a polished, editable Word document with identical content & badges:
identity, consensus anchor, subtype matrix, 10-column epi compare tables (year /
sample-size / patient-type first) with conflict notes, mortality, the red
"需人工核对清单", and a final reference list.

Same rules as SKILL.md: every value keeps its badge(s); red stays loud; numbers are
never averaged across studies; checklist always present; all source links collected
into the final 参考来源 (Reference list) section.

Palette: muted dark-blue / slate (low contrast) to match build_report.py.

Requires python-docx (`pip install python-docx`).
"""

import sys
import json
import argparse

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BADGE_LABEL = {
    "official": "\U0001F3DB官方", "db": "✅库", "guideline": "\U0001F4D8指南", "web": "\U0001F50Dweb",
    "weak": "⚠存疑", "redflag": "\U0001F534必核", "na": "❓N/A",
}
# muted dark-blue / slate family; redflag kept a desaturated clay so the critical flag is still legible
BADGE_HEX = {
    "official": "2B4A6F", "db": "35637A", "guideline": "4A5A82", "web": "5C748E",
    "weak": "8A7B5E", "redflag": "9D4B4B", "na": "94A3B8",
}
ALIASES = {"gov": "official", "registry": "official", "regulator": "official",
           "database": "db", "orphanet": "db", "pubmed": "db", "ctgov": "db",
           "guide": "guideline", "consensus": "guideline", "genereviews": "guideline",
           "media": "web", "pmid": "web", "secondary": "weak", "unknown": "na",
           "notfound": "na", "red": "redflag"}


def _norm(k):
    if not k:
        return None
    k = str(k).strip()
    return k if k in BADGE_LABEL else ALIASES.get(k.lower())


def _badge_keys(badge):
    if badge is None:
        return []
    if isinstance(badge, (list, tuple)):
        return [b for b in (_norm(x) for x in badge) if b]
    n = _norm(badge)
    return [n] if n else []


def add_badges(par, badge):
    for k in _badge_keys(badge):
        r = par.add_run("  " + BADGE_LABEL[k])
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(BADGE_HEX[k])


def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)


def set_widths(table, widths_in):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_in):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def add_hyperlink(par, url, text, color="2B4A6F"):
    part = par.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rPr.append(col)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    par._p.append(link)


def banner(doc, text, hexfill, hextext):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(hextext)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    pPr.append(shd)
    return p


def kv_section(doc, title, items):
    if not items:
        return
    doc.add_heading(title, level=1)
    for it in items:
        if isinstance(it, str):
            doc.add_paragraph(it, style="List Bullet")
            continue
        p = doc.add_paragraph(style="List Bullet")
        label = it.get("label") or it.get("field")
        text = it.get("text") or it.get("value") or ""
        if label:
            rb = p.add_run(str(label) + "：")
            rb.bold = True
        p.add_run(str(text))
        add_badges(p, it.get("badge"))


EPI_COLS = [
    ("year", "年份"), ("sample_size", "样本量/base"), ("patient_type", "患者类型/纳入"),
    ("metric", "metric"), ("value", "数值+CI"), ("geo_level", "地理"),
    ("denominator", "分母"), ("case_definition", "病例定义"), ("study_design", "设计"),
]

# muted palette shading tokens
FILL_HEADER = "EEF2F7"      # slate header
FILL_FIRSTCOL = "F1F5F9"    # row-label
FILL_HL = "E4EAF1"          # first-3 epi columns / emphasis (muted slate-blue)
FILL_CK = "EFE6E6"          # checklist header (muted clay tint)
TXT_CK = "7A4242"           # checklist header text (muted clay)


def main(argv):
    ap = argparse.ArgumentParser(description="Render a rare-disease-epi Word (.docx) report from JSON.")
    ap.add_argument("data")
    ap.add_argument("-o", "--out", default=None)
    args = ap.parse_args(argv)
    data = json.load(open(args.data, encoding="utf-8"))
    meta = data.get("meta", {})

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    disease = meta.get("disease") or meta.get("name") or "罕见病流行病学报告"
    doc.add_heading("罕见病流行病学报告：" + disease, level=0)
    sub_bits = []
    for k, pre in (("orphacode", ""), ("omim", "OMIM "), ("icd10", "ICD-10 "), ("icd11", "ICD-11 "), ("gard", "GARD ")):
        if meta.get(k):
            sub_bits.append(pre + str(meta[k]))
    if sub_bits:
        doc.add_paragraph("  ·  ".join(sub_bits))
    metabits = ["International \U0001F30D + 美国 \U0001F1FA\U0001F1F8 + 中国 \U0001F1E8\U0001F1F3",
                "生成 " + str(meta.get("generated", ""))]
    if meta.get("mode"):
        metabits.append("范围：" + str(meta["mode"]))
    mp = doc.add_paragraph("  ·  ".join(metabits))
    mp.runs[0].font.size = Pt(9)
    mp.runs[0].font.color.rgb = RGBColor.from_string("64748B")

    if meta.get("consensus_anchor"):
        banner(doc, "\U0001F4D8 共识锚点：" + str(meta["consensus_anchor"]), "EEF2F7", "2B4A6F")
    banner(doc, "⚠ 仅供研究情报参考，非医疗建议、非诊断。每个 epi 数字须带年份+样本量+徽标；标 \U0001F534 项必人工复核。",
           "EEF2F7", "3A4A63")
    if meta.get("skipped_gates"):
        banner(doc, "⚠ 已按用户要求一次跑完，未逐项停下确认；下面所有标 \U0001F534 的地方都还没经过人工核对。",
               "F6F0F0", "7A4242")

    # legend
    doc.add_heading("徽标说明（七档）", level=1)
    lp = doc.add_paragraph()
    for k in ("official", "db", "guideline", "web", "weak", "redflag", "na"):
        r = lp.add_run(BADGE_LABEL[k] + "  ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(BADGE_HEX[k])
    doc.add_paragraph("可信度：🏛️官方 ＝ ✅库 ＝ 📘指南 ＞ 🔍web ＞ ⚠ ＞ ❓N/A。🔴 是独立的'必核'标，"
                      "中国侧 / 死亡率 / 外推值即使来自官方/指南也叠 🔴。")

    # identity
    ident = data.get("identity") or []
    if ident:
        doc.add_heading("身份卡", level=1)
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        for it in ident:
            cells = t.add_row().cells
            cells[0].text = ""
            rr = cells[0].paragraphs[0].add_run(str(it.get("field", "")))
            rr.bold = True
            shade(cells[0], FILL_FIRSTCOL)
            cells[1].text = str(it.get("value", ""))
            add_badges(cells[1].paragraphs[0], it.get("badge"))
        set_widths(t, [2.1, 4.4])

    kv_section(doc, "共识锚点（病例定义 / 诊断标准 / 亚型分类 📘）", data.get("consensus_anchor"))

    # subtype matrix
    sm = data.get("subtype_matrix")
    if sm and sm.get("rows"):
        doc.add_heading("亚型矩阵（携带频率 vs 临床患病率分列）", level=1)
        cols = sm.get("columns") or ["亚型", "国际", "美国", "中国"]
        t = doc.add_table(rows=1, cols=len(cols))
        t.style = "Table Grid"
        for i, h in enumerate(cols):
            t.rows[0].cells[i].text = ""
            rp = t.rows[0].cells[i].paragraphs[0].add_run(h)
            rp.bold = True
            shade(t.rows[0].cells[i], FILL_HEADER)
        for r in sm["rows"]:
            cells = t.add_row().cells
            cells[0].text = ""
            rn = cells[0].paragraphs[0].add_run(str(r.get("subtype", "")))
            rn.bold = True
            shade(cells[0], FILL_FIRSTCOL)
            for idx, c in enumerate(r.get("cells") or [], start=1):
                if idx >= len(cells):
                    break
                c = c or {}
                cells[idx].text = str(c.get("value", "—"))
                if c.get("kind"):
                    sp = cells[idx].add_paragraph(str(c.get("kind")))
                    sp.runs[0].font.size = Pt(8.5)
                    sp.runs[0].font.color.rgb = RGBColor.from_string("64748B")
                add_badges(cells[idx].paragraphs[0], c.get("badge"))
        if sm.get("note"):
            np = doc.add_paragraph(str(sm["note"]))
            np.runs[0].font.size = Pt(8.5)
            np.runs[0].font.color.rgb = RGBColor.from_string("64748B")

    # epi compare tables (10 columns)
    for tbl in (data.get("epi_tables") or []):
        rows = tbl.get("rows") or []
        if not rows:
            continue
        doc.add_heading(str(tbl.get("title") or "epi 对比表") + "  （10 列口径 · 不取平均）", level=1)
        headers = [lbl for _, lbl in EPI_COLS] + ["来源+徽标"]
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = ""
            rp = t.rows[0].cells[i].paragraphs[0].add_run(h)
            rp.bold = True
            rp.font.size = Pt(8)
            shade(t.rows[0].cells[i], FILL_HL if i < 3 else FILL_HEADER)
        for r in rows:
            cells = t.add_row().cells
            for i, (key, _) in enumerate(EPI_COLS):
                cells[i].text = str(r.get(key, "") or "")
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(8)
            src = r.get("source") or {}
            sc = cells[len(EPI_COLS)]
            if isinstance(src, dict):
                sc.text = str(src.get("label", "") or "")
                if sc.paragraphs[0].runs:
                    sc.paragraphs[0].runs[0].font.size = Pt(8)
                if src.get("url"):
                    lp2 = sc.add_paragraph()
                    add_hyperlink(lp2, src["url"], src["url"][:48])
                add_badges(sc.paragraphs[0], src.get("badge"))
            else:
                sc.text = str(src)
        set_widths(t, [0.7, 0.95, 1.0, 0.7, 0.9, 0.6, 0.6, 0.95, 0.6, 1.0])
        if tbl.get("conflict_note"):
            cp = doc.add_paragraph()
            cr = cp.add_run("冲突解读：")
            cr.bold = True
            cp.add_run(str(tbl["conflict_note"]))
            cp.runs[0].font.color.rgb = RGBColor.from_string("3F4B5E")

    kv_section(doc, "死亡率（整项 🔴：罕见病死亡登记普遍不全）", data.get("mortality"))

    # red checklist
    ck = data.get("checklist") or []
    doc.add_heading("🔴 需人工核对清单（本工具最关键产出）", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    for i, h in enumerate(["#", "字段", "当前值 / 状态", "为什么要核", "建议去哪个源核"]):
        t.rows[0].cells[i].text = ""
        rp = t.rows[0].cells[i].paragraphs[0].add_run(h)
        rp.bold = True
        rp.font.color.rgb = RGBColor.from_string(TXT_CK)
        shade(t.rows[0].cells[i], FILL_CK)
    for i, it in enumerate(ck, 1):
        cells = t.add_row().cells
        cells[0].text = str(i)
        cells[1].text = str(it.get("field", ""))
        cells[2].text = str(it.get("value", ""))
        add_badges(cells[2].paragraphs[0], it.get("badge"))
        cells[3].text = str(it.get("why", ""))
        cells[4].text = str(it.get("where", ""))
    set_widths(t, [0.4, 1.7, 1.7, 1.7, 1.7])

    # references
    refs = data.get("references")
    if refs:
        doc.add_heading("参考来源（Reference list）", level=1)
        doc.add_paragraph("下列为本报告所有引用链接的汇总；请优先以 🏛️官方 / ✅库 / 📘指南 溯源核对。")
        for r in refs:
            if isinstance(r, str):
                url, label, badge = r, r, None
            else:
                url = r.get("url") or ""
                label = r.get("label") or url
                badge = r.get("badge")
            p = doc.add_paragraph(style="List Number")
            p.add_run(str(label) + "  ")
            if url:
                add_hyperlink(p, url, url)
            add_badges(p, badge)

    out = args.out or (args.data.rsplit(".", 1)[0] + ".docx")
    doc.save(out)
    print("Wrote " + out)
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
